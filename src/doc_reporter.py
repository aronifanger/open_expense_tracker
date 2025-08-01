"""
DOCX Reporter Module

This module is responsible for both fetching report data and rendering it
into a Word (.docx) document. This separation allows the same data to be
used for different outputs, like a Streamlit page.
"""
import logging
import pandas as pd
from datetime import datetime
from docx import Document
from src import config
from io import BytesIO

def get_report_data(processing_date: datetime, period: str) -> dict:
    """
    Fetches and compiles all necessary data for a report into a dictionary of DataFrames.
    """
    date_str = processing_date.strftime('%Y-%m-%d')
    file_suffix = f"{date_str}_{period}_deputy_scores.csv"
    deputy_scores_path = config.REPORTS_DIR / file_suffix
    critical_expenses_path = config.REPORTS_DIR / f"{date_str}_{period}_critical_expenses.csv"

    if not deputy_scores_path.exists() or not critical_expenses_path.exists():
        logging.warning(f"Report CSVs for period '{period}' not found.")
        return None

    deputy_scores_df = pd.read_csv(deputy_scores_path)
    critical_expenses_df = pd.read_csv(critical_expenses_path)
    
    top_10_deputies = deputy_scores_df.head(10)
    
    top_suppliers = pd.DataFrame()
    top_expense_types = pd.DataFrame()

    if not critical_expenses_df.empty:
        top_suppliers = critical_expenses_df['nomeFornecedor'].value_counts().nlargest(10).reset_index()
        top_suppliers.columns = ['Fornecedor', 'Nº de Ocorrências Críticas']
        
        top_expense_types = critical_expenses_df['tipoDespesa'].value_counts().nlargest(10).reset_index()
        top_expense_types.columns = ['Tipo de Despesa', 'Nº de Ocorrências Críticas']

    return {
        "top_10_deputies": top_10_deputies,
        "critical_expenses": critical_expenses_df,
        "top_suppliers": top_suppliers,
        "top_expense_types": top_expense_types
    }

def generate_word_report(processing_date: datetime, period: str, report_data: dict) -> BytesIO:
    """
    Generates the Word document from the pre-compiled report data and returns it as a byte stream.
    """
    logging.info(f"Generating Word report for date '{processing_date.strftime('%Y-%m-%d')}' and period '{period}'")
    if not report_data:
        logging.warning("Cannot generate Word report: report_data is empty.")
        return None

    doc = Document()
    date_str_formatted = processing_date.strftime('%d de %B de %Y')
    period_title = {'diário': 'Diário', 'semanal': 'Semanal', 'mensal': 'Mensal'}.get(period, period.capitalize())
    
    doc.add_heading(f'Relatório de Fiscalização {period_title} - {date_str_formatted}', level=1)

    # --- Summary Tables ---
    _add_summary_table(doc, report_data["top_10_deputies"], f'Top 10 Deputados com Despesas Críticas ({period_title})')
    _add_summary_table(doc, report_data["top_suppliers"], 'Top 10 Fornecedores em Despesas Críticas')
    _add_summary_table(doc, report_data["top_expense_types"], 'Top 10 Tipos de Despesa em Despesas Críticas')

    # --- Individual Analysis ---
    doc.add_heading('Análise Individual dos Top Deputados', level=1)
    for _, deputy in report_data["top_10_deputies"].iterrows():
        doc.add_heading(f"Deputado: {deputy['deputy_name']}", level=2)
        doc.add_paragraph("[Gráfico da evolução histórica será inserido aqui.]")
        
        deputy_expenses = report_data["critical_expenses"][
            report_data["critical_expenses"]['deputy_id'] == deputy['deputy_id']
        ].nlargest(1, 'score_fraude')
        
        if not deputy_expenses.empty:
            _add_summary_table(doc, deputy_expenses, f'Principal Despesa Crítica do Período ({period_title})')
        else:
            doc.add_paragraph("Nenhuma despesa crítica encontrada para este deputado no período.")
        doc.add_paragraph()
    
    # Save document to a byte stream
    doc_io = BytesIO()
    doc.save(doc_io)
    doc_io.seek(0)
    logging.info("Word report generated successfully and saved to memory stream.")
    return doc_io

def _add_summary_table(doc, df, title):
    """Adds a formatted table to the Word document."""
    if title:
        doc.add_heading(title, level=2)
    if df.empty:
        doc.add_paragraph("Nenhum dado encontrado.")
        return
    
    table = doc.add_table(rows=1, cols=len(df.columns))
    table.style = 'Table Grid'
    for i, col_name in enumerate(df.columns):
        table.cell(0, i).text = str(col_name)
    for _, row in df.iterrows():
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = str(val)
